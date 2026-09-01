import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}>'
                      f'<w:top w:w="{top}" w:type="dxa"/>'
                      f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
                      f'<w:left w:w="{left}" w:type="dxa"/>'
                      f'<w:right w:w="{right}" w:type="dxa"/>'
                      f'</w:tcMar>')
    tcPr.append(tcMar)

def create_report():
    doc = docx.Document()
    
    # Page setup - 1 inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Styles & Colors
    PRIMARY_COLOR = RGBColor(27, 54, 93)      # Deep Navy Blue
    SECONDARY_COLOR = RGBColor(41, 128, 185)  # Soft Blue
    TEXT_COLOR = RGBColor(44, 62, 80)         # Charcoal Dark
    MUTED_BG = "F4F6F9"
    HEADER_BG = "1B365D"
    
    # Document Title
    title_p = doc.add_paragraph()
    title_run = title_p.add_run("Pond Catchment Analysis Backend\nTechnical Report")
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = PRIMARY_COLOR
    title_p.paragraph_format.space_after = Pt(6)
    title_p.paragraph_format.space_before = Pt(0)
    
    subtitle_p = doc.add_paragraph()
    sub_run = subtitle_p.add_run("Terrain Analysis, Hydrological D8 Routing, and Automated Pond Planning")
    sub_run.font.size = Pt(12)
    sub_run.font.italic = True
    sub_run.font.color.rgb = SECONDARY_COLOR
    subtitle_p.paragraph_format.space_after = Pt(18)
    
    # Metadata Box / Summary Table
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    meta_info = [
        ("GitHub Repository", "https://github.com/nikhilesh-git/pond-planning-system"),
        ("Working API Endpoint (Remote)", "http://10.1.75.51:3000/analyzeContour"),
        ("Local Development URL", "http://127.0.0.1:3000/analyzeContour"),
        ("API Framework & Protocol", "FastAPI / Uvicorn (HTTP POST multipart/form-data)")
    ]
    
    for i, (k, v) in enumerate(meta_info):
        cell_k = table.cell(i, 0)
        cell_v = table.cell(i, 1)
        cell_k.width = Inches(2.2)
        cell_v.width = Inches(4.3)
        
        pk = cell_k.paragraphs[0]
        rk = pk.add_run(k)
        rk.font.bold = True
        rk.font.size = Pt(10)
        rk.font.color.rgb = PRIMARY_COLOR
        
        pv = cell_v.paragraphs[0]
        rv = pv.add_run(v)
        rv.font.size = Pt(10)
        if "http" in v:
            rv.font.color.rgb = SECONDARY_COLOR
            rv.font.underline = True
            
        set_cell_background(cell_k, "EBF2FA")
        set_cell_background(cell_v, "FAFCFF")
        set_cell_margins(cell_k, top=120, bottom=120, left=150, right=150)
        set_cell_margins(cell_v, top=120, bottom=120, left=150, right=150)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    
    # Section 1: Working API Route URL
    h1 = doc.add_heading(level=1)
    r1 = h1.add_run("1. Working API Route & Deployment")
    r1.font.color.rgb = PRIMARY_COLOR
    
    p = doc.add_paragraph()
    p.add_run("The backend API service is deployed on the remote server and actively listening on port 3000:")
    
    bullet1 = doc.add_paragraph(style='List Bullet')
    bullet1.add_run("Remote Server Base URL: ").font.bold = True
    bullet1.add_run("http://10.1.75.51:3000")
    
    bullet2 = doc.add_paragraph(style='List Bullet')
    bullet2.add_run("Analyze Contour Route: ").font.bold = True
    bullet2.add_run("http://10.1.75.51:3000/analyzeContour")
    
    bullet3 = doc.add_paragraph(style='List Bullet')
    bullet3.add_run("Interactive API Documentation (Swagger UI): ").font.bold = True
    bullet3.add_run("http://10.1.75.51:3000/docs")

    # Section 2: Catchment Estimation Approach
    h2 = doc.add_heading(level=1)
    r2 = h2.add_run("2. Catchment Estimation Approach")
    r2.font.color.rgb = PRIMARY_COLOR
    
    p = doc.add_paragraph()
    p.add_run("The system performs a fully dynamic, multi-stage geospatial and hydrological pipeline to analyze terrain directly from raw contour data without relying on hardcoded coordinates.")
    
    steps = [
        ("KML/KMZ Parsing", "The uploaded contour file is parsed using high-speed 'lxml.etree'. LineStrings and their associated elevation values (Z attributes) are extracted with zero manual intervention."),
        ("Coordinate Projection (UTM Transformation)", "Using 'pyproj', the algorithm automatically detects the optimal local Universal Transverse Mercator (UTM) zone from the bounding box of coordinates. Geographic coordinates (WGS84 Latitude/Longitude) are projected into metric Cartesian units (meters), ensuring high precision for distance, slope, and catchment area calculations."),
        ("Digital Elevation Model (DEM) Generation", "Contour point data is rasterized into a high-resolution 2-meter regular grid. SciPy's 'griddata' (Linear interpolation with nearest-neighbor extrapolation) constructs a continuous 2D surface elevation matrix z(x,y)."),
        ("Hydrological Routing (D8 Flow Algorithm)", "The D8 (Deterministic 8-neighbor) flow direction model determines the path of water descent across every cell. Flow accumulation is computed to simulate how surface runoff converges downstream."),
        ("Identification of Natural Sinks (Depressions)", "Local minima where water naturally gathers without immediate outward flow paths are marked as candidate sinks. For each candidate sink, upstream tracing determines the full contributing catchment area."),
        ("Multi-Stage Waterbody & River Exclusion Pipeline", "To prevent selecting invalid locations (e.g., inside existing rivers, floodplains, or edge artifacts), candidate sinks pass through rigorous filtering:\n"
         "  • Boundary Buffer Exclusion: Rejects sinks within 200m of the DEM edges.\n"
         "  • Elevation Band Filter: Excludes sinks in the bottom 10% elevation range.\n"
         "  • Flow Accumulation Filter: Sinks sitting on or directly adjacent to major river channels are rejected.\n"
         "  • Flat Waterbody Rejection: Sinks within extensive flat zones (>1200 sqm) are flagged as existing lakes/reservoirs and excluded."),
        ("Multi-Factor Suitability Scoring", "Remaining candidate sinks are scored using a weighted suitability index:\n"
         "  Total Score = (0.30 × Elev_Score) + (0.40 × Catchment_Score) + (0.30 × Accum_Score)\n"
         "  - Elevation Score favors upland sites to enable gravity irrigation and avoid riverbeds.\n"
         "  - Catchment Score rewards moderate, balanced catchments to prevent overflow risk.\n"
         "  - Accumulation Score favors true micro-depressions over high-volume drainage lines.")
    ]
    
    for title, desc in steps:
        p = doc.add_paragraph()
        r_t = p.add_run(f"• {title}: ")
        r_t.font.bold = True
        r_t.font.color.rgb = PRIMARY_COLOR
        p.add_run(desc)
        p.paragraph_format.space_after = Pt(4)

    # Section 3: Demonstration using Provided Contour Map
    h3 = doc.add_heading(level=1)
    r3 = h3.add_run("3. Demonstration on Provided Contour Map")
    r3.font.color.rgb = PRIMARY_COLOR
    
    p = doc.add_paragraph()
    p.add_run("The API was validated by submitting the standard dataset (")
    p.add_run("contours_1m.kml").font.bold = True
    p.add_run(") to the ")
    p.add_run("/analyzeContour").font.bold = True
    p.add_run(" endpoint. The results obtained are as follows:")
    
    # Results Table
    res_table = doc.add_table(rows=4, cols=2)
    res_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    res_table.autofit = False
    
    results_data = [
        ("Optimal Latitude", "21.25205° N"),
        ("Optimal Longitude", "81.30924° E"),
        ("Sink Elevation", "290.0 meters"),
        ("Estimated Catchment Area", "772.0 – 780.0 m²")
    ]
    
    for i, (k, v) in enumerate(results_data):
        c0 = res_table.cell(i, 0)
        c1 = res_table.cell(i, 1)
        c0.width = Inches(2.5)
        c1.width = Inches(4.0)
        
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(k)
        r0.font.bold = True
        r0.font.size = Pt(10)
        
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(v)
        r1.font.bold = True
        r1.font.size = Pt(10)
        r1.font.color.rgb = PRIMARY_COLOR
        
        set_cell_background(c0, "F0F4F8")
        set_cell_background(c1, "FFFFFF")
        set_cell_margins(c0, top=100, bottom=100, left=150, right=150)
        set_cell_margins(c1, top=100, bottom=100, left=150, right=150)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    
    p_analysis = doc.add_paragraph()
    p_analysis.add_run("Hydrological Evaluation of Selected Site:\n").font.bold = True
    p_analysis.add_run("The selected site sits at an elevation of 290.0m on an upland micro-depression. It successfully avoids the lower elevation river corridor (286m) that runs across the area. The catchment area provides sufficient inflow without risking structural embankment failure or heavy silt deposition.")

    # Section 4: API Documentation
    h4 = doc.add_heading(level=1)
    r4 = h4.add_run("4. Complete API Documentation")
    r4.font.color.rgb = PRIMARY_COLOR
    
    # Endpoint spec
    doc.add_paragraph().add_run("POST /analyzeContour").font.bold = True
    doc.add_paragraph("Uploads a KML or KMZ contour map and returns the computed optimal pond location along with its catchment area.")
    
    doc.add_heading(level=2).add_run("Request Specifications")
    bullet_req1 = doc.add_paragraph(style='List Bullet')
    bullet_req1.add_run("HTTP Method: ").font.bold = True
    bullet_req1.add_run("POST")
    
    bullet_req2 = doc.add_paragraph(style='List Bullet')
    bullet_req2.add_run("Content-Type: ").font.bold = True
    bullet_req2.add_run("multipart/form-data")
    
    bullet_req3 = doc.add_paragraph(style='List Bullet')
    bullet_req3.add_run("Body Parameters: ").font.bold = True
    bullet_req3.add_run("file (Binary .kml or .kmz file, Required)")

    doc.add_heading(level=2).add_run("Response Specification")
    bullet_res1 = doc.add_paragraph(style='List Bullet')
    bullet_res1.add_run("Status Code: ").font.bold = True
    bullet_res1.add_run("200 OK")
    
    bullet_res2 = doc.add_paragraph(style='List Bullet')
    bullet_res2.add_run("Content-Type: ").font.bold = True
    bullet_res2.add_run("application/json")
    
    doc.add_paragraph("Example JSON Response Body:").runs[0].font.italic = True
    
    code_box = doc.add_table(rows=1, cols=1)
    code_box.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_code = code_box.cell(0, 0)
    cell_code.width = Inches(6.5)
    set_cell_background(cell_code, "F8F9FA")
    set_cell_margins(cell_code, top=140, bottom=140, left=180, right=180)
    
    code_text = (
        "{\n"
        '  "pond_location": {\n'
        '    "latitude": 21.252048140122685,\n'
        '    "longitude": 81.30923688499112,\n'
        '    "elevation": 290.0\n'
        '  },\n'
        '  "catchment_area_sqm": 772.0\n'
        "}"
    )
    p_code = cell_code.paragraphs[0]
    r_code = p_code.add_run(code_text)
    r_code.font.name = "Consolas"
    r_code.font.size = Pt(9.5)
    r_code.font.color.rgb = RGBColor(30, 30, 30)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    
    doc.add_heading(level=2).add_run("Example cURL Command")
    curl_box = doc.add_table(rows=1, cols=1)
    curl_box.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_curl = curl_box.cell(0, 0)
    c_curl.width = Inches(6.5)
    set_cell_background(c_curl, "F8F9FA")
    set_cell_margins(c_curl, top=140, bottom=140, left=180, right=180)
    
    curl_text = (
        'curl -X POST "http://10.1.75.51:3000/analyzeContour" \\\n'
        '     -H "accept: application/json" \\\n'
        '     -H "Content-Type: multipart/form-data" \\\n'
        '     -F "file=@contours_1m.kml"'
    )
    p_curl = c_curl.paragraphs[0]
    r_curl = p_curl.add_run(curl_text)
    r_curl.font.name = "Consolas"
    r_curl.font.size = Pt(9.5)
    r_curl.font.color.rgb = RGBColor(30, 30, 30)

    output_path = "Pond_Catchment_Analysis_Report.docx"
    doc.save(output_path)
    print(f"Successfully generated {output_path}")

if __name__ == "__main__":
    create_report()
